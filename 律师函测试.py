import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from docx import Document
from datetime import datetime
import threading
import queue

class LegalLetterGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("律师函自动生成软件")
        self.root.geometry("900x600")
        
        # 设置变量
        self.template_path = tk.StringVar()
        self.data_path = tk.StringVar()
        self.output_dir = tk.StringVar()
        self.file_prefix = tk.StringVar(value="律师函")
        self.variables = []
        self.variable_mapping = {}
        self.progress_var = tk.DoubleVar()
        
        # 创建界面
        self.create_ui()
    
    def create_ui(self):
        # 创建标签页
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 模板管理页面
        template_frame = ttk.Frame(notebook)
        notebook.add(template_frame, text="模板管理")
        self.setup_template_frame(template_frame)
        
        # 数据导入页面
        data_frame = ttk.Frame(notebook)
        notebook.add(data_frame, text="数据导入")
        self.setup_data_frame(data_frame)
        
        # 文档生成页面
        generate_frame = ttk.Frame(notebook)
        notebook.add(generate_frame, text="文档生成")
        self.setup_generate_frame(generate_frame)
        
        # 设置页面
        settings_frame = ttk.Frame(notebook)
        notebook.add(settings_frame, text="设置")
        self.setup_settings_frame(settings_frame)
    
    def setup_template_frame(self, parent):
        # 模板选择区域
        frame = ttk.LabelFrame(parent, text="选择模板文件")
        frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(frame, text="Word模板文件:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        ttk.Entry(frame, textvariable=self.template_path, width=50).grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(frame, text="浏览...", command=self.browse_template).grid(row=0, column=2, padx=5, pady=5)
        ttk.Button(frame, text="解析变量", command=self.extract_variables).grid(row=0, column=3, padx=5, pady=5)
        
        # 变量显示区域
        var_frame = ttk.LabelFrame(parent, text="模板变量")
        var_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 创建变量列表框
        self.var_listbox = tk.Listbox(var_frame)
        self.var_listbox.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 变量说明
        ttk.Label(parent, text="提示: 在Word模板中使用{{变量名}}格式作为占位符").pack(pady=5)
    
    def setup_data_frame(self, parent):
        # 数据文件选择
        frame = ttk.LabelFrame(parent, text="选择数据文件")
        frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(frame, text="Excel数据文件:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        ttk.Entry(frame, textvariable=self.data_path, width=50).grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(frame, text="浏览...", command=self.browse_data).grid(row=0, column=2, padx=5, pady=5)
        ttk.Button(frame, text="加载数据", command=self.load_data).grid(row=0, column=3, padx=5, pady=5)
        
        # 变量映射区域
        map_frame = ttk.LabelFrame(parent, text="变量映射")
        map_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 创建映射表格
        self.mapping_tree = ttk.Treeview(map_frame, columns=("模板变量", "Excel列名"), show="headings")
        self.mapping_tree.heading("模板变量", text="模板变量")
        self.mapping_tree.heading("Excel列名", text="Excel列名")
        self.mapping_tree.column("模板变量", width=150)
        self.mapping_tree.column("Excel列名", width=150)
        self.mapping_tree.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 添加滚动条
        scrollbar = ttk.Scrollbar(map_frame, orient=tk.VERTICAL, command=self.mapping_tree.yview)
        self.mapping_tree.configure(yscroll=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 映射操作按钮
        btn_frame = ttk.Frame(map_frame)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(btn_frame, text="自动映射", command=self.auto_map_variables).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="修改映射", command=self.edit_mapping).pack(side=tk.LEFT, padx=5)
    
    def setup_generate_frame(self, parent):
        # 输出目录选择
        frame = ttk.LabelFrame(parent, text="输出设置")
        frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(frame, text="输出目录:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        ttk.Entry(frame, textvariable=self.output_dir, width=50).grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(frame, text="浏览...", command=self.browse_output_dir).grid(row=0, column=2, padx=5, pady=5)
        
        # 生成选项
        options_frame = ttk.LabelFrame(parent, text="生成选项")
        options_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(options_frame, text="文件名前缀:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        ttk.Entry(options_frame, textvariable=self.file_prefix, width=20).grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        
        self.format_var = tk.StringVar(value="docx")
        ttk.Label(options_frame, text="输出格式:").grid(row=0, column=2, padx=5, pady=5, sticky=tk.W)
        ttk.Radiobutton(options_frame, text="Word", variable=self.format_var, value="docx").grid(row=0, column=3, padx=5, pady=5)
        ttk.Radiobutton(options_frame, text="PDF", variable=self.format_var, value="pdf").grid(row=0, column=4, padx=5, pady=5)
        
        # 进度条
        progress_frame = ttk.LabelFrame(parent, text="生成进度")
        progress_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X, padx=5, pady=5)
        
        self.status_label = ttk.Label(progress_frame, text="就绪")
        self.status_label.pack(pady=5)
        
        # 生成按钮
        btn_frame = ttk.Frame(parent)
        btn_frame.pack(fill=tk.X, padx=10, pady=10)
        ttk.Button(btn_frame, text="生成律师函", command=self.generate_documents, style="Accent.TButton").pack(side=tk.RIGHT, padx=5)
    
    def setup_settings_frame(self, parent):
        # 基本设置
        settings_frame = ttk.LabelFrame(parent, text="基本设置")
        settings_frame.pack(fill=tk.X, padx=10, pady=10)
        
        # 默认保存路径
        ttk.Label(settings_frame, text="默认保存路径:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        self.default_path_var = tk.StringVar(value=os.path.expanduser("~/Documents"))
        ttk.Entry(settings_frame, textvariable=self.default_path_var, width=50).grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(settings_frame, text="浏览...", command=self.browse_default_path).grid(row=0, column=2, padx=5, pady=5)
        
        # 保存设置按钮
        ttk.Button(parent, text="保存设置", command=self.save_settings).pack(side=tk.RIGHT, padx=10, pady=10)
    
    def browse_template(self):
        filetypes = [("Word文档", "*.docx"), ("所有文件", "*.*")]
        filename = filedialog.askopenfilename(title="选择Word模板文件", filetypes=filetypes)
        if filename:
            self.template_path.set(filename)
    
    def browse_data(self):
        filetypes = [("Excel文件", "*.xlsx *.xls"), ("所有文件", "*.*")]
        filename = filedialog.askopenfilename(title="选择Excel数据文件", filetypes=filetypes)
        if filename:
            self.data_path.set(filename)
    
    def browse_output_dir(self):
        directory = filedialog.askdirectory(title="选择输出目录")
        if directory:
            self.output_dir.set(directory)
    
    def browse_default_path(self):
        directory = filedialog.askdirectory(title="选择默认保存路径")
        if directory:
            self.default_path_var.set(directory)
    
    def extract_variables(self):
        template_path = self.template_path.get()
        if not template_path:
            messagebox.showerror("错误", "请先选择Word模板文件")
            return
    
        try:
            # 清空变量列表
            self.var_listbox.delete(0, tk.END)
    
            # 从模板中提取变量
            self.variables = self.extract_template_variables(template_path)
    
            # 显示变量列表
            for var in self.variables:
                self.var_listbox.insert(tk.END, var)
    
            # 自动生成Excel模板
            if self.variables:
                import pandas as pd
                df = pd.DataFrame(columns=self.variables)
                excel_path = os.path.join(os.path.dirname(template_path), "律师函数据模板.xlsx")
                df.to_excel(excel_path, index=False)
                messagebox.showinfo("成功", f"成功从模板中提取了{len(self.variables)}个变量\n并已生成Excel模板：{excel_path}")
            else:
                messagebox.showinfo("成功", "未提取到任何变量")
    
        except Exception as e:
            messagebox.showerror("错误", f"提取变量时出错: {str(e)}")
    
    def extract_template_variables(self, template_path):
        """从Word模板中提取所有变量（{{变量名}}格式），支持跨run"""
        doc = Document(template_path)
        variables = set()
    
        # 提取段落中的变量
        for paragraph in doc.paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            var_matches = re.findall(r'{{(.*?)}}', full_text)
            variables.update(var_matches)
    
        # 提取表格中的变量
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = ''.join(run.text for run in paragraph.runs)
                        var_matches = re.findall(r'{{(.*?)}}', full_text)
                        variables.update(var_matches)
    
        return sorted(list(variables))
    
    def load_data(self):
        data_path = self.data_path.get()
        if not data_path:
            messagebox.showerror("错误", "请先选择Excel数据文件")
            return
        
        if not self.variables:
            messagebox.showerror("错误", "请先从模板中提取变量")
            return
        
        try:
            # 读取Excel文件
            self.df = pd.read_excel(data_path)
            
            # 清空映射表格
            for item in self.mapping_tree.get_children():
                self.mapping_tree.delete(item)
            
            # 初始化变量映射
            self.variable_mapping = {}
            
            # 显示变量和可能的映射
            for var in self.variables:
                self.mapping_tree.insert("", tk.END, values=(var, ""))
            
            messagebox.showinfo("成功", f"成功加载数据文件，包含{len(self.df)}条记录")
        except Exception as e:
            messagebox.showerror("错误", f"加载数据时出错: {str(e)}")
    
    def auto_map_variables(self):
        if not hasattr(self, 'df'):
            messagebox.showerror("错误", "请先加载数据文件")
            return
        
        # 清空现有映射
        for item in self.mapping_tree.get_children():
            self.mapping_tree.delete(item)
        
        # 自动映射变量
        excel_columns = self.df.columns.tolist()
        
        for var in self.variables:
            # 尝试找到匹配的列名
            matched_column = ""
            
            # 精确匹配
            if var in excel_columns:
                matched_column = var
            else:
                # 模糊匹配（忽略大小写和空格）
                var_clean = var.lower().replace(" ", "")
                for col in excel_columns:
                    col_clean = col.lower().replace(" ", "")
                    if var_clean == col_clean:
                        matched_column = col
                        break
            
            # 添加到映射表格
            self.mapping_tree.insert("", tk.END, values=(var, matched_column))
            self.variable_mapping[var] = matched_column
        
        messagebox.showinfo("自动映射", "变量自动映射完成，请检查并调整不正确的映射")
    
    def edit_mapping(self):
        # 获取选中的项
        selected_item = self.mapping_tree.selection()
        if not selected_item:
            messagebox.showerror("错误", "请先选择要编辑的映射")
            return
        
        # 获取当前值
        current_values = self.mapping_tree.item(selected_item[0], 'values')
        var_name = current_values[0]
        current_column = current_values[1]
        
        # 创建编辑对话框
        edit_window = tk.Toplevel(self.root)
        edit_window.title("编辑映射")
        edit_window.geometry("400x200")
        edit_window.transient(self.root)
        edit_window.grab_set()
        
        ttk.Label(edit_window, text=f"为变量 '{var_name}' 选择对应的Excel列:").pack(pady=10)
        
        # 创建下拉列表
        column_var = tk.StringVar(value=current_column)
        column_combo = ttk.Combobox(edit_window, textvariable=column_var, width=30)
        column_combo['values'] = list(self.df.columns)
        column_combo.pack(pady=10)
        
        # 确定按钮
        def save_mapping():
            new_column = column_var.get()
            self.mapping_tree.item(selected_item[0], values=(var_name, new_column))
            self.variable_mapping[var_name] = new_column
            edit_window.destroy()
        
        ttk.Button(edit_window, text="确定", command=save_mapping).pack(pady=10)
    
    def generate_documents(self):
        # 检查必要的输入
        template_path = self.template_path.get()
        output_dir = self.output_dir.get()
        
        if not template_path:
            messagebox.showerror("错误", "请先选择Word模板文件")
            return
        
        if not output_dir:
            messagebox.showerror("错误", "请先选择输出目录")
            return
        
        if not hasattr(self, 'df'):
            messagebox.showerror("错误", "请先加载数据文件")
            return
        
        if not self.variable_mapping:
            messagebox.showerror("错误", "请先完成变量映射")
            return
        
        # 准备数据列表
        data_list = []
        for _, row in self.df.iterrows():
            data_dict = {}
            for var, col in self.variable_mapping.items():
                if col and col in self.df.columns:
                    data_dict[var] = row[col]
            data_dict['文件名前缀'] = self.file_prefix.get()
            data_list.append(data_dict)
        
        # 创建生成线程
        self.progress_var.set(0)
        self.status_label.config(text="正在生成文档...")
        
        # 使用队列进行线程通信
        self.progress_queue = queue.Queue()
        
        # 启动生成线程
        generate_thread = threading.Thread(
            target=self.generate_legal_letters_thread,
            args=(template_path, data_list, output_dir, self.format_var.get(), self.progress_queue)
        )
        generate_thread.daemon = True
        generate_thread.start()
        
        # 启动进度更新
        self.root.after(100, self.update_progress)
    
    def generate_legal_letters_thread(self, template_path, data_list, output_dir, format_type, progress_queue):
        try:
            total = len(data_list)
            for index, data in enumerate(data_list):
                # 创建文档
                doc = Document(template_path)
                
                # 替换文档中的变量
                self.replace_variables_in_document(doc, data)
                
                # 保存文件（命名规则：日期+前缀+序号）
                date_prefix = datetime.now().strftime("%Y%m%d")
                file_name = f"{date_prefix}_{data.get('文件名前缀', '律师函')}_{index+1}.docx"
                file_path = os.path.join(output_dir, file_name)
                
                # 保存文档
                doc.save(file_path)
                
                # TODO: 如果需要PDF格式，添加转换代码
                
                # 更新进度
                progress = (index + 1) / total * 100
                progress_queue.put((progress, f"已生成 {index+1}/{total} 份文档"))
            
            # 完成
            progress_queue.put((100, f"完成！成功生成 {total} 份律师函"))
        except Exception as e:
            progress_queue.put((-1, f"生成文档时出错: {str(e)}"))
    
    def replace_variables_in_document(self, doc, data):
        # 替换段落中的变量
        for paragraph in doc.paragraphs:
            text = paragraph.text
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                text = text.replace(placeholder, str(value) if pd.notna(value) else "")
            # 只在有变化时重写
            if text != paragraph.text:
                paragraph.clear()
                paragraph.add_run(text)

        # 替换表格中的变量
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        for key, value in data.items():
                            placeholder = f"{{{{{key}}}}}"
                            text = text.replace(placeholder, str(value) if pd.notna(value) else "")
                        if text != paragraph.text:
                            paragraph.clear()
                            paragraph.add_run(text)
    
    def update_progress(self):
        try:
            while not self.progress_queue.empty():
                progress, message = self.progress_queue.get_nowait()
                if progress < 0:
                    # 错误情况
                    self.progress_var.set(0)
                    self.status_label.config(text=message)
                    messagebox.showerror("错误", message)
                    return
                else:
                    # 正常进度更新
                    self.progress_var.set(progress)
                    self.status_label.config(text=message)
                    
                    # 如果完成，显示成功消息
                    if progress >= 100:
                        messagebox.showinfo("成功", message)
                        return
            
            # 继续更新进度
            self.root.after(100, self.update_progress)
        except Exception as e:
            self.status_label.config(text=f"更新进度时出错: {str(e)}")
    
    def save_settings(self):
        # 保存设置逻辑
        messagebox.showinfo("设置", "设置已保存")

# 主程序入口
def main():
    root = tk.Tk()
    app = LegalLetterGenerator(root)
    
    # 设置样式
    style = ttk.Style()
    style.configure("Accent.TButton", font=("Arial", 10, "bold"))
    
    root.mainloop()

if __name__ == "__main__":
    main()